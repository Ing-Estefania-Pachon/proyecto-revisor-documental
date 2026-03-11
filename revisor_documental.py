import docx
import pandas as pd
import google.generativeai as genai
import json
import os

# CONFIGURACIÓN: Reemplaza con tu API KEY de Gemini
genai.configure(api_key="TU_API_KEY_AQUÍ")
model = genai.GenerativeModel('gemini-1.5-flash')

def ejecutar_skill_fase1(texto_documento):
    # Prompt Maestro basado en Metodología v2 [cite: 370, 541]
    prompt_fase1 = f"""
    Actúa como un Auditor Editorial Senior de la UNGRD[cite: 541]. 
    Realiza la Fase 1 de Revisión de Estilo (Guía de Autor)[cite: 409].
    
    CRITERIOS A EVALUAR:
    1. Título: Max 30 palabras[cite: 420].
    2. Resumen: Max 300 palabras[cite: 423].
    3. Palabras Clave: 5 a 6 términos[cite: 424].
    4. Estructura: Debe tener Título, Autores, Afiliaciones, Contacto, Resumen, Palabras clave y Bibliografía[cite: 419, 422].

    INSTRUCCIÓN DE SALIDA: Responde ÚNICAMENTE en formato JSON estricto con esta estructura:
    {{
      "resultados": [
        {{"criterio": "Título", "estado": "Cumple", "hallazgo": "X palabras", "recomendacion": "Ninguna"}},
        {{"criterio": "Resumen", "estado": "No Cumple", "hallazgo": "X palabras", "recomendacion": "Recortar"}}
      ],
      "estado_final": "APROBADO/RECHAZADO"
    }}

    TEXTO DEL DOCUMENTO:
    {texto_documento}
    """
    
    response = model.generate_content(prompt_fase1)
    # Limpiamos la respuesta para obtener solo el JSON
    json_data = response.text.replace('```json', '').replace('```', '').strip()
    return json.loads(json_data)

def generar_reportes_automaticos(nombre_archivo):
    if not os.path.exists(nombre_archivo):
        print(f"❌ No se encuentra el archivo: {nombre_archivo}")
        return

    print(f"🤖 El Agente ARES está revisando la Fase 1 de: {nombre_archivo}...")
    doc = docx.Document(nombre_archivo)
    
    # Extraemos solo el inicio del documento (Metadatos) para no gastar tokens innecesarios
    texto_inicio = "\n".join([p.text for p in doc.paragraphs[:15]])
    
    # Llamada al agente
    resultado_ia = ejecutar_skill_fase1(texto_inicio)
    
    # --- GENERAR EXCEL (Tabla de Cumplimiento) ---
    df = pd.DataFrame(resultado_ia["resultados"])
    df.to_excel("Reporte_Cumplimiento_ARES.xlsx", index=False)
    
    # --- GENERAR WORD (Recomendaciones) ---
    reporte_doc = docx.Document()
    reporte_doc.add_heading('Reporte de Revisión de Estilo - Fase 1', 0)
    reporte_doc.add_paragraph(f"Estado Final: {resultado_ia['estado_final']}")
    
    for res in resultado_ia["resultados"]:
        if res["estado"] == "No Cumple":
            p = reporte_doc.add_paragraph()
            p.add_run(f"🔴 {res['criterio']}: ").bold = True
            p.add_run(f"{res['hallazgo']}. Recomendación: {res['recomendacion']}")
            
    reporte_doc.save("Reporte_Recomendaciones_Fase1.docx")
    print(f"✅ Revisión finalizada. Estado: {resultado_ia['estado_final']}")

if __name__ == "__main__":
    generar_reportes_automaticos("FICHA_CORDOBA.docx")